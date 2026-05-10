"""PDF/DOCX text extraction with SHA-256 cache.

Three extractor paths, dispatched on file extension and content:

  1. ``.docx`` → ``python-docx``                    [DOCX_MODEL_VERSION]
     Sejm prints often ship a clean editable .docx alongside the signed-PDF
     scan. CLI prefers .docx when present (cli._resolve_pdf_relpath). Perfect
     Polish text, no OCR. ~0.1s/file.

  2. ``.pdf`` w/ text layer → ``pymupdf4llm``       [PYMUPDF_MODEL_VERSION]
     Digital Sejm prints carry typed text. ``PymupdfBackend`` first probes
     pages with plain ``pymupdf`` to detect scans, then runs
     ``pymupdf4llm.to_markdown`` on text-bearing pages only — yielding
     **structured markdown** (## headings, **bold**, lists). pymupdf4llm has
     internal Tesseract auto-OCR but defaults to English which destroys
     Polish diacritics; we sidestep that by feeding only text-layer pages.
     ~0.5s/page, no GPU.

  3. ``.pdf`` scan (0 chars) → ``tesseract`` + pol  [TESSERACT_MODEL_VERSION]
     Scanned signed letters / opinions / transmittals. ~140/543 prints in
     the corpus. PymupdfBackend returns empty result; ``extract_pdf``
     auto-routes to ``_extract_tesseract`` which rasterizes via pymupdf at
     ``TESSERACT_DPI`` and OCRs each page with Polish (``pol``) traineddata.
     ~3-5s/page CPU, plain text output (no markdown — scans are mostly
     short transmittal letters where heading structure is minimal anyway).
     LightOnOCR-1B / paddle-style transformer OCR was rejected: would
     repeat the paddle GPU/RAM disaster (multi-GB VRAM, deadlock risk on
     Windows + RTX 5060 Ti) for marginal markdown gain on short scans.

  4. ``.pdf`` w/ ``SUPAGRAF_PDF_BACKEND=paddle`` → ``PaddleOCRVL``
     Legacy escape hatch — PaddleOCR-VL-1.5 with long-document slicing
     (head 8 + tail 2) still wired behind env flag. NOT default.

Cache keyed by ``(sha256, model_version)``: version bumps add rows, never
overwrite, so old entries stay for audit and new format versions trigger
fresh extraction.
"""
from __future__ import annotations

import concurrent.futures
import hashlib
import os
import tempfile
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Protocol

from loguru import logger
from postgrest.exceptions import APIError
from pypdf import PdfReader, PdfWriter
from tenacity import retry, retry_if_exception_type, stop_after_attempt, wait_exponential

from supagraf.db import supabase

# pymupdf version pinned at runtime via fitz; cache version string includes
# the package version so a future upgrade doesn't quietly serve stale cache.
try:
    import fitz  # type: ignore[import-not-found]
    _PYMUPDF_VERSION = getattr(fitz, "__version__", "unknown")
except Exception:
    fitz = None  # type: ignore[assignment]
    _PYMUPDF_VERSION = "unavailable"

PYMUPDF_MODEL_VERSION = f"pymupdf-{_PYMUPDF_VERSION}-md-primary"
PADDLE_MODEL_VERSION = "paddle-vl-1.5-primary"
PADDLE_MODEL_VERSION_SLICED = "paddle-vl-1.5-primary-sliced"
PYPDF_FALLBACK_VERSION = "pypdf-6.10.2-fallback"
DOCX_MODEL_VERSION = "python-docx-1.2-primary"
TESSERACT_MODEL_VERSION = "tesseract-5.5-pol-fallback"

# Tesseract OCR — auto-fallback when pymupdf returns 0 chars (scanned PDF).
# DPI for rasterizing pages — 300 is the sweet spot for typed Polish text;
# higher hurts speed without measurably improving accuracy on Sejm scans.
TESSERACT_DPI = int(os.environ.get("SUPAGRAF_TESSERACT_DPI", "300"))
TESSERACT_LANG = os.environ.get("SUPAGRAF_TESSERACT_LANG", "pol")
TESSERACT_ENABLED = os.environ.get("SUPAGRAF_TESSERACT_ENABLED", "1") not in ("0", "false", "")

# Slicing thresholds — applied only to paddle backend (pymupdf is fast enough
# to extract full documents).
SLICE_THRESHOLD_PAGES = 10  # strict ``>`` — 10 stays full
SLICE_HEAD_PAGES = 8
SLICE_TAIL_PAGES = 2

# Paddle deadlock watchdog — ~1-in-5 PDFs deadlock paddle silently. Per-call
# timeout fires PaddleTimeoutError; caller routes to pypdf fallback.
PADDLE_TIMEOUT_S = float(os.environ.get("SUPAGRAF_PADDLE_TIMEOUT_S", "180"))

# Default backend selector. "pymupdf" (default) | "paddle".
PDF_BACKEND = os.environ.get("SUPAGRAF_PDF_BACKEND", "pymupdf").lower()


class PaddleTimeoutError(RuntimeError):
    """Raised when paddle ``predict`` exceeds ``PADDLE_TIMEOUT_S``."""


def _paddle_predict_with_timeout(model: Any, pdf_path: Path) -> list[Any]:
    ex = concurrent.futures.ThreadPoolExecutor(
        max_workers=1, thread_name_prefix="paddle-predict"
    )
    try:
        fut = ex.submit(lambda: list(model.predict(str(pdf_path))))
        try:
            return fut.result(timeout=PADDLE_TIMEOUT_S)
        except concurrent.futures.TimeoutError as e:
            raise PaddleTimeoutError(
                f"paddle deadlock on {pdf_path.name} after {PADDLE_TIMEOUT_S}s"
            ) from e
    finally:
        ex.shutdown(wait=False)


def _should_slice(page_count: int) -> bool:
    return page_count > SLICE_THRESHOLD_PAGES


def _slice_pages(pages: list, page_count: int) -> tuple[list, str]:
    if not _should_slice(page_count):
        return pages, "full"
    head = pages[:SLICE_HEAD_PAGES]
    tail = pages[-SLICE_TAIL_PAGES:]
    return head + tail, "first8_last2"


def _build_sliced_pdf(src: Path, dst: Path, page_count: int) -> int:
    reader = PdfReader(str(src))
    indexes = list(range(min(SLICE_HEAD_PAGES, page_count)))
    tail_start = max(page_count - SLICE_TAIL_PAGES, len(indexes))
    indexes.extend(range(tail_start, page_count))
    writer = PdfWriter()
    for i in indexes:
        writer.add_page(reader.pages[i])
    with dst.open("wb") as f:
        writer.write(f)
    return len(indexes)


def _make_slice_header(page_count: int, pages_extracted: int) -> str:
    return (
        f"<!-- extracted_pages={pages_extracted} strategy=first8_last2 "
        f"source_pages={page_count} -->\n"
    )


def _make_slice_omission_marker(page_count: int) -> str:
    omit_start = SLICE_HEAD_PAGES + 1
    omit_end = page_count - SLICE_TAIL_PAGES
    return (
        f"\n\n<!-- pages {omit_start}..{omit_end} omitted; first {SLICE_HEAD_PAGES} "
        f"+ last {SLICE_TAIL_PAGES} only (long-document strategy "
        f"{PADDLE_MODEL_VERSION_SLICED}) -->\n\n"
    )


@dataclass(frozen=True)
class ExtractionResult:
    sha256: str
    text: str
    page_count: int
    ocr_used: bool
    char_count_per_page: list[int]
    model_version: str
    cache_hit: bool


class OcrBackend(Protocol):
    def extract(self, pdf_path: Path) -> tuple[str, list[int]]:
        """Return (full_text, char_count_per_page)."""


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _extract_pypdf(path: Path) -> tuple[str, list[int]]:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for p in reader.pages:
        pages.append(p.extract_text() or "")
    return "\n\n".join(pages), [len(t) for t in pages]


def _extract_tesseract(path: Path) -> tuple[str, list[int]]:
    """OCR a scanned PDF with Tesseract.

    Uses pymupdf to rasterize each page to PIL image at TESSERACT_DPI, then
    pytesseract for OCR with Polish language data. No poppler dependency.

    Tesseract v5.x + ``pol`` traineddata must be installed on PATH; verified
    at extraction time. ~5-10s per page on CPU.
    """
    if fitz is None:
        raise RuntimeError("pymupdf required for rasterization — `uv add pymupdf`")
    try:
        import pytesseract
        from PIL import Image
        import io
    except ImportError as e:
        raise RuntimeError("pytesseract + Pillow required for OCR fallback") from e

    pages_text: list[str] = []
    per_page: list[int] = []
    zoom = TESSERACT_DPI / 72.0  # pymupdf default is 72 DPI; scale up for OCR
    matrix = fitz.Matrix(zoom, zoom)
    with fitz.open(path) as doc:
        for page in doc:
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            page_text = pytesseract.image_to_string(img, lang=TESSERACT_LANG).strip()
            pages_text.append(page_text)
            per_page.append(len(page_text))
    if sum(per_page) == 0:
        raise RuntimeError(
            f"tesseract OCR returned 0 chars across {len(per_page)} pages for {path.name}"
        )
    return "\n\n".join(pages_text), per_page


# Signer-section keywords. The retry-on-miss heuristic only fires when NONE
# of these appear in the cover text, so they must be SPECIFIC to the
# wnioskodawcy block — generic terms like "projekt" or "Druk nr" appear on
# every Sejm cover and would suppress every retry.
_COVER_HEADER_KEYWORDS = (
    "wnioskodawc",                    # "Wnioskodawcy" / "wnioskodawców"
    "podpisał",                       # "Podpisał..." (occasional spelling)
    "podpisali",
    "wniosek poselski wnoszą",        # boiler standard for poselski projekt
    "poselski projekt ustawy wnoszą",
    "(-)",                             # Polish signature notation "(-) Jan Kowalski"
)


def extract_pdf_cover(pdf_path: Path, max_pages: int = 2) -> str:
    """First N pages plain text — used to recover MP signers list on Sejm
    'Poselski projekt' PDFs.

    Sygnatariusze (15-30 nazwisk) drukowani są zawsze na pierwszej-drugiej
    stronie (occasionally pushed to page 3-4 by long preamble). Docx body
    je gubi (CLI prefers .docx over .pdf in ``_resolve_pdf_relpath``). Plain
    text wystarczy — LLM wyciąga listę.

    No cache: cover extraction is ~100ms per PDF, not worth a separate
    pdf_extracts row keyed by max_pages. Returns empty string if pymupdf
    unavailable or file missing.
    """
    if fitz is None:
        raise RuntimeError("pymupdf required for cover extraction — `uv add pymupdf`")
    if not pdf_path.exists():
        return ""
    parts: list[str] = []
    with fitz.open(pdf_path) as doc:
        total_pages = doc.page_count
        cap = min(max_pages, total_pages)
        for i in range(cap):
            parts.append(doc.load_page(i).get_text("text"))
    text = "\n\n".join(parts).strip()
    if not text:
        return text

    # v3 (citizen-review issue #3): if requested depth missed the signers
    # block but the document is longer than what we read, push deeper. Uses
    # heuristic header keywords; bump max_pages to 4 once on miss to catch
    # covers where a long preamble pushes "Wnioskodawcy" past page 2.
    lowered = text.lower()
    if max_pages < 4 and not any(kw.lower() in lowered for kw in _COVER_HEADER_KEYWORDS):
        with fitz.open(pdf_path) as doc:
            total_pages = doc.page_count
        if total_pages > max_pages:
            return extract_pdf_cover(pdf_path, max_pages=4)
    return text


def _extract_docx(path: Path) -> tuple[str, list[int]]:
    """Extract text from a .docx (Office Open XML) file.

    Sejm prints often ship as both a signed-PDF (scanned) and a clean .docx
    (the editable original). The .docx side gives perfect text without OCR.
    """
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("python-docx not installed — run `uv add python-docx`") from e
    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text]
    # Tables — flatten cell-by-cell, preserve row breaks.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    text = "\n\n".join(paragraphs).strip()
    if not text:
        raise RuntimeError(f"docx returned 0 chars for {path.name}")
    # docx has no concept of pages; report a single "page" for the cache shape.
    return text, [len(text)]


class PymupdfBackend:
    """Markdown extraction via PyMuPDF + pymupdf4llm.

    Two-step strategy:
      1. Plain pymupdf text extraction per page — fast, deterministic, used to
         detect scanned (no text layer) pages without burning OCR cycles.
      2. If any page has text, run ``pymupdf4llm.to_markdown`` on the full
         document to produce structured markdown (headings, lists, bold).
         pymupdf4llm has internal Tesseract OCR auto-fallback but it defaults
         to English which destroys Polish diacritics — we steer around it by
         feeding only pages that already have a text layer.

    Pages with 0 chars are reported as such; ``extract_pdf`` then routes the
    document to ``_extract_tesseract`` with ``pol`` language for proper Polish
    OCR.
    """

    def extract(self, pdf_path: Path) -> tuple[str, list[int]]:
        if fitz is None:
            raise RuntimeError("pymupdf not installed — run `uv add pymupdf`")
        # Step 1: per-page plain-text probe to detect scans without OCR.
        per_page: list[int] = []
        with fitz.open(pdf_path) as doc:
            for page in doc:
                per_page.append(len(page.get_text("text").strip()))

        if sum(per_page) == 0:
            # All pages scan-only — let extract_pdf route to tesseract+pol.
            return "", per_page

        # Step 2: structured markdown via pymupdf4llm. Restrict to pages that
        # have text so its English-only Tesseract auto-OCR isn't triggered.
        try:
            import pymupdf4llm  # type: ignore[import-not-found]
        except ImportError:
            # Fall back to plain text if pymupdf4llm absent — still usable.
            with fitz.open(pdf_path) as doc:
                texts = [p.get_text("text").strip() for p in doc]
            return "\n\n".join(texts), per_page

        text_pages = [i for i, n in enumerate(per_page) if n > 0]
        md = pymupdf4llm.to_markdown(
            str(pdf_path),
            pages=text_pages,
            show_progress=False,
        )
        return md, per_page


class PaddleOcrBackend:
    """PaddleOCR-VL-1.5 backend (heavy, GPU-bound). Lazy singleton.

    Only used when ``SUPAGRAF_PDF_BACKEND=paddle`` or as explicit fallback.
    """

    _model: Any = None
    _model_lock = threading.Lock()

    @classmethod
    def _get_model(cls) -> Any:
        if cls._model is not None:
            return cls._model
        with cls._model_lock:
            if cls._model is not None:
                return cls._model
            try:
                from paddleocr import PaddleOCRVL
            except ImportError as e:
                raise ImportError(
                    "PaddleOCR not installed. Run: uv add paddleocr paddlepaddle-gpu"
                ) from e
            device = os.environ.get("SUPAGRAF_PADDLE_DEVICE")
            if not device:
                try:
                    import paddle  # type: ignore[import]
                    device = "gpu" if paddle.is_compiled_with_cuda() else "cpu"
                except Exception:
                    device = "cpu"
            logger.info("loading PaddleOCR-VL-1.5 on device={} (cold start)", device)
            cls._model = PaddleOCRVL(pipeline_version="v1.5", device=device)
            return cls._model

    def extract(self, pdf_path: Path) -> tuple[str, list[int]]:
        model = self._get_model()
        results = _paddle_predict_with_timeout(model, pdf_path)
        if not results:
            raise RuntimeError(f"PaddleOCR returned no pages for {pdf_path.name}")
        pages_text: list[str] = []
        per_page: list[int] = []
        for res in results:
            page_text = ""
            md = getattr(res, "markdown", None)
            if isinstance(md, dict):
                page_text = (md.get("markdown_texts") or "").strip()
            if not page_text:
                blocks = (res.json.get("res", {}) or {}).get("parsing_res_list", []) if hasattr(res, "json") else []
                page_text = "\n".join(
                    str(b.get("block_content", "")) for b in blocks if isinstance(b, dict)
                ).strip()
            pages_text.append(page_text)
            per_page.append(len(page_text))
        if sum(per_page) == 0:
            raise RuntimeError(
                f"PaddleOCR returned 0 chars across {len(per_page)} pages for {pdf_path.name}"
            )
        return "\n\n".join(pages_text), per_page


def _resolve_default_backend() -> tuple[OcrBackend, str]:
    """Return (backend, model_version) for ``SUPAGRAF_PDF_BACKEND`` env."""
    if PDF_BACKEND == "pymupdf":
        return PymupdfBackend(), PYMUPDF_MODEL_VERSION
    if PDF_BACKEND == "paddle":
        return PaddleOcrBackend(), PADDLE_MODEL_VERSION
    raise NotImplementedError(f"unknown SUPAGRAF_PDF_BACKEND={PDF_BACKEND!r}")


@retry(
    retry=retry_if_exception_type(APIError),
    stop=stop_after_attempt(4),
    wait=wait_exponential(multiplier=1, min=1, max=8),
    reraise=True,
)
def _cache_lookup(sha: str, model_version: str) -> dict | None:
    r = (
        supabase()
        .table("pdf_extracts")
        .select("text, page_count, ocr_used, char_count_per_page")
        .eq("sha256", sha)
        .eq("model_version", model_version)
        .limit(1)
        .execute()
    )
    return (r.data or [None])[0]


@retry(
    retry=retry_if_exception_type(APIError),
    stop=stop_after_attempt(4),
    wait=wait_exponential(multiplier=1, min=1, max=8),
    reraise=True,
)
def _cache_insert(sha: str, model_version: str, text: str, page_count: int,
                  ocr_used: bool, per_page: list[int]) -> None:
    supabase().table("pdf_extracts").upsert(
        {
            "sha256": sha,
            "model_version": model_version,
            "text": text,
            "page_count": page_count,
            "ocr_used": ocr_used,
            "char_count_per_page": per_page,
        },
        on_conflict="sha256,model_version",
    ).execute()


def _count_pdf_pages(path: Path) -> int:
    return len(PdfReader(str(path)).pages)


def _extract_paddle_sliced(
    backend: OcrBackend, src: Path, page_count: int
) -> tuple[str, list[int]]:
    tf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tf.close()
    sliced_path = Path(tf.name)
    try:
        pages_written = _build_sliced_pdf(src, sliced_path, page_count)
        text, per_page = backend.extract(sliced_path)
        if len(per_page) != pages_written:
            logger.warning(
                "paddle returned {} pages from sliced pdf but {} were written for {}",
                len(per_page), pages_written, src.name,
            )
        header = _make_slice_header(page_count, len(per_page))
        omission = _make_slice_omission_marker(page_count)
        chunks = text.split("\n\n")
        if len(chunks) == len(per_page) and len(per_page) > SLICE_HEAD_PAGES:
            head_chunks = chunks[:SLICE_HEAD_PAGES]
            tail_chunks = chunks[SLICE_HEAD_PAGES:]
            new_text = (
                header
                + "\n\n".join(head_chunks)
                + omission
                + "\n\n".join(tail_chunks)
            )
        else:
            new_text = header + text + omission
        return new_text, per_page
    finally:
        try:
            sliced_path.unlink()
        except OSError:
            pass


def extract_pdf(path: Path, *, ocr_backend: OcrBackend | None = None) -> ExtractionResult:
    """Extract text. Default backend = pymupdf (fast, no GPU).

    Dispatches by file extension:
      - .docx → python-docx (clean text, no OCR)
      - .pdf  → pymupdf (or paddle via ``SUPAGRAF_PDF_BACKEND=paddle``)

    Override with ``SUPAGRAF_PDF_BACKEND=paddle`` for layout-aware OCR. Slicing
    only applies to paddle. ``ocr_backend`` parameter overrides env (used by
    tests).
    """
    if not path.exists():
        raise FileNotFoundError(path)

    # docx fast-path — no OCR backend involved.
    if path.suffix.lower() == ".docx":
        sha = _sha256(path)
        cached = _cache_lookup(sha, DOCX_MODEL_VERSION)
        if cached is not None:
            return ExtractionResult(
                sha, cached["text"], cached["page_count"], cached["ocr_used"],
                cached["char_count_per_page"], DOCX_MODEL_VERSION, cache_hit=True,
            )
        text, per_page = _extract_docx(path)
        _cache_insert(sha, DOCX_MODEL_VERSION, text, len(per_page), False, per_page)
        return ExtractionResult(
            sha, text, len(per_page), False, per_page, DOCX_MODEL_VERSION, cache_hit=False
        )

    sha = _sha256(path)

    # Resolve backend + model_version. Tests can pass ocr_backend explicitly.
    if ocr_backend is None:
        backend, model_version = _resolve_default_backend()
        is_paddle = isinstance(backend, PaddleOcrBackend)
    else:
        backend = ocr_backend
        is_paddle = isinstance(backend, PaddleOcrBackend)
        model_version = PADDLE_MODEL_VERSION if is_paddle else PYMUPDF_MODEL_VERSION

    # Slicing only meaningful for paddle. pymupdf handles full documents fast.
    sliced = False
    page_count = 0
    if is_paddle:
        try:
            page_count = _count_pdf_pages(path)
        except Exception as e:
            logger.warning("could not count pages for {}: {!r} — full extract", path.name, e)
            page_count = 0
        sliced = _should_slice(page_count)
        if sliced:
            model_version = PADDLE_MODEL_VERSION_SLICED

    cached = _cache_lookup(sha, model_version)
    if cached is not None:
        return ExtractionResult(
            sha,
            cached["text"],
            cached["page_count"],
            cached["ocr_used"],
            cached["char_count_per_page"],
            model_version,
            cache_hit=True,
        )

    primary_exc: Exception | None = None
    try:
        if is_paddle and sliced:
            text, per_page = _extract_paddle_sliced(backend, path, page_count)
        else:
            text, per_page = backend.extract(path)
        if sum(per_page) == 0:
            # pymupdf returns 0 chars on scanned PDFs (no text layer). Auto-OCR
            # with tesseract before falling back to pypdf — pypdf will also
            # return 0 chars on scans, so without OCR these prints are lost.
            if not is_paddle and TESSERACT_ENABLED:
                logger.info(
                    "pymupdf 0 chars for {} — trying tesseract OCR fallback",
                    path.name,
                )
                cached = _cache_lookup(sha, TESSERACT_MODEL_VERSION)
                if cached is not None:
                    return ExtractionResult(
                        sha, cached["text"], cached["page_count"], cached["ocr_used"],
                        cached["char_count_per_page"], TESSERACT_MODEL_VERSION,
                        cache_hit=True,
                    )
                text, per_page = _extract_tesseract(path)
                _cache_insert(sha, TESSERACT_MODEL_VERSION, text, len(per_page), True, per_page)
                return ExtractionResult(
                    sha, text, len(per_page), True, per_page, TESSERACT_MODEL_VERSION,
                    cache_hit=False,
                )
            raise RuntimeError(
                f"backend returned 0 chars across {len(per_page)} pages for {path.name}"
            )
        if is_paddle and sliced:
            expected = SLICE_HEAD_PAGES + SLICE_TAIL_PAGES
            if len(per_page) != expected:
                logger.warning(
                    "paddle returned {} pages for {} but expected {} (sliced)",
                    len(per_page), path.name, expected,
                )
        ocr_used = is_paddle  # pymupdf is text extraction, not OCR
        _cache_insert(sha, model_version, text, len(per_page), ocr_used, per_page)
        return ExtractionResult(
            sha, text, len(per_page), ocr_used, per_page, model_version, cache_hit=False
        )
    except (KeyboardInterrupt, SystemExit):
        raise
    except Exception as e:
        primary_exc = e
        logger.warning(
            "primary backend ({}) failed for {}: {!r} — pypdf fallback",
            type(backend).__name__, path.name, e,
        )

    try:
        text, per_page = _extract_pypdf(path)
    except (KeyboardInterrupt, SystemExit):
        raise
    except Exception as pypdf_exc:
        logger.error("pypdf fallback also failed for {}: {!r}", path.name, pypdf_exc)
        assert primary_exc is not None
        raise primary_exc

    if sum(per_page) == 0:
        raise RuntimeError(
            f"both primary and pypdf produced 0 chars for {path.name} "
            f"(primary error: {primary_exc!r})"
        )

    _cache_insert(sha, PYPDF_FALLBACK_VERSION, text, len(per_page), False, per_page)
    return ExtractionResult(
        sha, text, len(per_page), False, per_page, PYPDF_FALLBACK_VERSION, cache_hit=False
    )
